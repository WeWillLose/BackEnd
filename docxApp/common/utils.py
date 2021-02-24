import csv
import io
import zipfile
import pandas as pd
from docx import Document
from docx.shared import Inches
import numpy as np


def read_docx_tables(file, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        file:    Word Document string path or file

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(file)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise


def zip_files(files):
    """
    collect dict of files to in memory zip archive (.zip)

    Parameters:
        files:     dict of files.
                    When [None] - return empty bytes

    Note: if file in files is empty, it will be skipped

    Return: bytes of zip archive with files
    """

    outfile = io.BytesIO()

    with zipfile.ZipFile(outfile, 'w') as zf:
        if files is None:
            pass
        else:
            for n, f in enumerate(files):
                tmp = files.get(f)
                if tmp:
                    zf.writestr("{}".format(tmp.name), tmp.file.getvalue())
    return outfile.getvalue()


def do_report(data):
    """
    create report from input data

    Params:
        data:  list ...
    """
    report = Document("template.docx")
    data_ind = 0
    print(data)
    print("--------------")
    for table in report.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == '<!ball>':
                    tem = cell.text
                    print(data[data_ind])
                    print()
                    if data[data_ind] is None:
                        cell.text = cell.text.replace('<!ball>', '-')
                    else:
                        t = cell
                        cell.text = cell.text.replace('<!ball>', str(data[data_ind].sum()))
                        cell.paragraphs[0].paragraph_format.first_line_indent = Inches(0)
                        if data_ind == len(data)-1:
                            data_ind =0
                        else:
                            data_ind +=1
    report.save("report.docx")


def pars_table(table):
    """
    Parsing table and transform str to float

    Raise:
        ValueError when table has no column Баллы

    Params:
        table: table with column Баллы

    Return column Баллы
    """
    if "Баллы" in table:
        for ind, v in enumerate(table["Баллы"]):
            if type(v) == str:
                v = float(v.replace(',', '.'))
                table["Баллы"][ind] = v
        return table["Баллы"]
    raise ValueError("Table has no column - \"Баллы\"")


def pars_tables(tables):
    if tables is None or len(tables) == 0:
        raise Exception("Empty tables")
    res = []
    for table in tables:
        if "Баллы" in table:
            res.append(pars_table(table))
    return res


def pars_doc(file):
    if file is not None:
        tables = read_docx_tables(file)
    else:
        tables = read_docx_tables("data.docx")

    if len(tables) != 31:
        raise Exception("Присутсвуют не все таблицы в файле")

    parsed_data = pars_tables(tables)
    report = do_report(parsed_data)
